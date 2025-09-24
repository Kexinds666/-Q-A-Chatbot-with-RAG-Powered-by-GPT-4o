"""
Document processing service with advanced chunking strategies
"""

import os
import hashlib
import time
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import aiofiles
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from app.core.config import settings
from app.core.logging import logger, log_document_processing
from app.core.database import Document, DocumentChunk, get_db
from sqlalchemy.ext.asyncio import AsyncSession


class DocumentService:
    """Advanced document processing service"""
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.processed_dir = Path(settings.PROCESSED_DIR)
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.batch_size = settings.BATCH_SIZE
        
        # Create directories if they don't exist
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize tokenizer for token counting
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception:
            self.tokenizer = None
    
    async def process_document(
        self,
        file_path: str,
        filename: str,
        user_id: int,
        db: AsyncSession
    ) -> Dict[str, Any]:
        """Process uploaded document with advanced chunking"""
        
        start_time = time.time()
        document_id = None
        
        try:
            # Calculate file hash
            file_hash = await self._calculate_file_hash(file_path)
            
            # Check if document already exists
            existing_doc = await self._get_document_by_hash(file_hash, db)
            if existing_doc:
                return {
                    "document_id": existing_doc.id,
                    "status": "already_exists",
                    "message": "Document already processed"
                }
            
            # Create document record
            document = await self._create_document_record(
                file_path, filename, file_hash, user_id, db
            )
            document_id = document.id
            
            # Extract text based on file type
            text_content = await self._extract_text(file_path, filename)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from document")
            
            # Advanced text chunking
            chunks = await self._chunk_text(text_content, filename)
            
            # Save chunks to database
            await self._save_chunks(chunks, document_id, db)
            
            # Update document status
            processing_time = time.time() - start_time
            await self._update_document_status(
                document_id, "processed", chunk_count=len(chunks),
                processing_time=processing_time, db=db
            )
            
            log_document_processing(
                document_id=document_id,
                filename=filename,
                status="success",
                processing_time=processing_time,
                chunk_count=len(chunks)
            )
            
            return {
                "document_id": document_id,
                "status": "success",
                "chunk_count": len(chunks),
                "processing_time": processing_time
            }
            
        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = str(e)
            
            if document_id:
                await self._update_document_status(
                    document_id, "failed", error_message=error_msg, db=db
                )
            
            log_document_processing(
                document_id=document_id or 0,
                filename=filename,
                status="failed",
                processing_time=processing_time,
                error=error_msg
            )
            
            raise
    
    async def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        async with aiofiles.open(file_path, "rb") as f:
            async for chunk in f:
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    async def _get_document_by_hash(self, file_hash: str, db: AsyncSession) -> Optional[Document]:
        """Check if document with same hash already exists"""
        from sqlalchemy import select
        result = await db.execute(
            select(Document).where(Document.content_hash == file_hash)
        )
        return result.scalar_one_or_none()
    
    async def _create_document_record(
        self, file_path: str, filename: str, file_hash: str, user_id: int, db: AsyncSession
    ) -> Document:
        """Create document record in database"""
        file_size = os.path.getsize(file_path)
        file_type = Path(filename).suffix.lower()
        
        document = Document(
            filename=filename,
            original_filename=filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file_type,
            content_hash=file_hash,
            status="processing",
            owner_id=user_id
        )
        
        db.add(document)
        await db.commit()
        await db.refresh(document)
        return document
    
    async def _extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from various file formats"""
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._extract_pdf_text(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await self._extract_docx_text(file_path)
        elif file_ext == '.txt':
            return await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as f:
                    return await f.read()
            except Exception as e:
                raise ValueError(f"Failed to extract text from TXT: {e}")
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {e}")
    
    async def _chunk_text(self, text: str, filename: str) -> List[Dict[str, Any]]:
        """Advanced text chunking with metadata"""
        # Create langchain document
        doc = LangchainDocument(page_content=text, metadata={"source": filename})
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Process chunks and add metadata
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            content = chunk.page_content.strip()
            if not content:
                continue
            
            # Calculate token count
            token_count = self._count_tokens(content)
            
            # Create chunk hash
            chunk_hash = hashlib.sha256(content.encode()).hexdigest()
            
            processed_chunks.append({
                "chunk_index": i,
                "content": content,
                "content_hash": chunk_hash,
                "token_count": token_count,
                "metadata": chunk.metadata
            })
        
        return processed_chunks
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        if self.tokenizer:
            return len(self.tokenizer.encode(text))
        else:
            # Fallback: rough estimation (4 characters per token)
            return len(text) // 4
    
    async def _save_chunks(
        self, chunks: List[Dict[str, Any]], document_id: int, db: AsyncSession
    ):
        """Save document chunks to database"""
        for chunk_data in chunks:
            chunk = DocumentChunk(
                document_id=document_id,
                chunk_index=chunk_data["chunk_index"],
                content=chunk_data["content"],
                content_hash=chunk_data["content_hash"],
                token_count=chunk_data["token_count"]
            )
            db.add(chunk)
        
        await db.commit()
    
    async def _update_document_status(
        self, document_id: int, status: str, chunk_count: int = None,
        processing_time: float = None, error_message: str = None, db: AsyncSession = None
    ):
        """Update document status in database"""
        from sqlalchemy import update
        
        update_data = {"status": status}
        if chunk_count is not None:
            update_data["chunk_count"] = chunk_count
        if processing_time is not None:
            update_data["processing_time"] = processing_time
        if error_message is not None:
            update_data["error_message"] = error_message
        
        await db.execute(
            update(Document).where(Document.id == document_id).values(**update_data)
        )
        await db.commit()
    
    async def get_document_chunks(
        self, document_id: int, db: AsyncSession
    ) -> List[DocumentChunk]:
        """Get all chunks for a document"""
        from sqlalchemy import select
        
        result = await db.execute(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.chunk_index)
        )
        return result.scalars().all()
    
    async def delete_document(self, document_id: int, user_id: int, db: AsyncSession) -> bool:
        """Delete document and all its chunks"""
        from sqlalchemy import select, delete
        
        # Get document
        result = await db.execute(
            select(Document).where(
                Document.id == document_id,
                Document.owner_id == user_id
            )
        )
        document = result.scalar_one_or_none()
        
        if not document:
            return False
        
        # Delete file
        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception as e:
            logger.warning(f"Failed to delete file {document.file_path}: {e}")
        
        # Delete document (chunks will be deleted by cascade)
        await db.execute(delete(Document).where(Document.id == document_id))
        await db.commit()
        
        return True
